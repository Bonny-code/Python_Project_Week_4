import requests
# we import requests which to get data from an API

import docx
# WE IMPORT DOCX to be able to display data in a word document

import random
# random is imported to use and pick any random item.

from pprint import pprint
# this will print the output in a nice format

# PRINTING TO A word document
park_guide = docx.Document()
park_guide.add_paragraph('National Park Travel Guide', 'Heading 1')

parK_list_url = 'https://national-parks-1150.azurewebsites.net/api/list'

parK_list_response = requests.get(parK_list_url).json()

print(parK_list_response)

five_random_parks = random.sample(parK_list_response, 5)
print(five_random_parks)

# https://national-parks-1150.azurewebsites.net/api/list/CAVE
# below will print out a park name with the code.
for park in five_random_parks:
    print(park)
    base_parK_details_url = 'https://national-parks-1150.azurewebsites.net/api/list'
    park_code = park['park_code']
    park_name = park['name']

    # this below is supposed to create a word document.
    park_guide.add_paragraph(park_name)
    print(base_parK_details_url)
    print(park_code)
    parK_detail_url = base_parK_details_url + park_code
    print(parK_detail_url)

    park_detail_response = requests.get(parK_detail_url).json()
    pprint(park_detail_response)


park_guide.save('ParkGuide.docx')