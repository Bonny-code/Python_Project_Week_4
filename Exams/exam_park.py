# we import requests which is a library we use to get data using an API
import requests
# We import random because we use random to get a list of random numbers or things
import random
# Docx is used to save the output into a Word document.
import docx

# we use import inches from docx which will be used to customise our image
# in the Word document.
from docx.shared import Inches

# we import plotly which is used to create a map in word document
# using longitude and latitude
# pandas is also used alongside plotly to create the map.
import plotly.express as px
import pandas as pd

# Here we are starting to decide what the Word document will look like
# we are adding a Heading to the page which will be displayed in the beginning
# we use heading 1 which is major heading
# park_guide = docx.Document()  ********
# park_guide.add_paragraph('National Park Travel Guide', 'Heading 1')  *****

# 1. Here we are saving the API into a variable which we can use as we code along
park_list_url = 'https://national-parks-1150.azurewebsites.net/api/list'

# we use the response keyword to get the response fom the API
# We then print the response.
response_list_of_all_parks = requests.get(park_list_url).json()
print(response_list_of_all_parks)

# 2. pick 5 random parks
# Here we are going to pick 5 random parks and display their information
# we use the random key word to pick random parks and make sure it is not repeating.
# we are saving the response in the list_of_random_parks_chosen variable
list_of_random_parks_chosen = random.sample(response_list_of_all_parks, 5)
print(list_of_random_parks_chosen)

# 3. Make a Word document.
park_word_document = docx.Document()
park_guide = docx.Document()
# the code below is going to act as the title of the page
park_word_document.add_paragraph('National Park Travel Guide', 'Title')

# 4 For each Park
# here we will use a for loop to print out certain information about each park
# such as the park code, URL
for one_park in list_of_random_parks_chosen:
    print(one_park)
    park_code = one_park['park_code']
    print(park_code)

    # We will concatenate API and the park code the get those information for that particular park
    # then print the response
    url_for_one_park_details = 'https://national-parks-1150.azurewebsites.net/api/' + park_code
    print(url_for_one_park_details)
    one_park_details_response = requests.get(url_for_one_park_details).json()
    print(one_park_details_response)

    # This is to get the park name for one park
    # We have to look for the key 'Name' in the API response to get park name
    park_name = one_park_details_response['name']
    print(park_name)
    # send the output of the park name to a Word document with the name as the page title
    park_word_document.add_paragraph(park_name, 'Heading 1')

    # This is to get the park weather
    # this will create a heading for weather and also a
    # paragraph describing the weather conditions of the park.
    park_weather = one_park_details_response['weather_overview']
    park_word_document.add_paragraph('Weather', 'Heading 1')
    park_word_document.add_paragraph(park_weather)

    # this is to get the description from the details of the park
    # we will add it to a Word document as using heading 1
    # the description will be a new paragraph.
    park_description = one_park_details_response['description']
    park_word_document.add_paragraph('Description', 'Heading 1')
    park_word_document.add_paragraph(park_description)

    # This is to get the park operating hours
    # we will create a hading and output the operating hours
    # into a new paragraph.
    park_functional_hours = one_park_details_response['park_operating_hours']
    park_word_document.add_paragraph('Operating_hours', 'Heading 1')
    park_word_document.add_paragraph(park_functional_hours)

    # this is for the park activities which will be output into
    # to a new paragraph in the Word document
    # the activities will be a list of strings
    park_word_document.add_paragraph('Activities', 'Heading 1')
    park_activities_list = one_park_details_response['activities']

    for activity in park_activities_list:
        # we use the style='List Bullet' to indicate what style we want which in this case is
        # bullet list
        # park_word_document.add_paragraph(style='List Bullet')
        park_word_document.add_paragraph(activity, style='List Bullet')

    # print(one_park_details_response.keys())
    # dict_keys(['activities', 'contact_info', 'description', 'location', 'name', 'one_park_images', 'Park_code',
    #            'park_operating_hours', 'Weather_overview'])
    contact_info_dictionary = one_park_details_response['contact_info']
    # print(contact_info_dictionary)
    # print(contact_info_dictionary.keys())

    # **Adding the contact information.
    park_word_document.add_paragraph('Contact Information', 'Heading 1')

    # getting the address
    address_text = contact_info_dictionary['address']
    print(address_text)
    # adding a header to a document
    park_word_document.add_paragraph('Address', 'Heading 2')
    park_word_document.add_paragraph(address_text)

    # Getting the website for the park
    # we will pull the url from the api and write it to a Word document
    park_website_url = contact_info_dictionary['url']
    print(park_website_url)
    park_word_document.add_paragraph('Website', 'Heading 2')  # This will add a new paragraph
    park_word_document.add_paragraph(park_website_url)

    # How to get longitude and latitude
    # we will start with the longitude
    # they are in a list called location so, we have to go into the location to
    # be able to get the values we want
    # adding a new paragraph called location with style of heading 1
    location_dictionary = one_park_details_response['location']
    park_word_document.add_paragraph('Location', 'Heading 1')

    park_longitude = location_dictionary['longitude']  # getting the longitude
    print(park_longitude)
    park_word_document.add_paragraph('longitude', 'Heading 2')  # adding the longitude to a new paragraph
    park_word_document.add_paragraph(park_longitude)

    # to get latitude
    # same as the longitude
    park_latitude = location_dictionary['latitude']  # creating a new variable called park longitude
    print(park_latitude)
    park_word_document.add_paragraph('latitude', 'Heading 2')
    park_word_document.add_paragraph(park_latitude)  # adding the values to a Word document

    # 5. How to get the Image
    park_word_document.add_paragraph('Park Images', 'Heading 1')
    image_data_list = one_park_details_response['nps_park_images']
    for image_data_dictionary in image_data_list:
        print(image_data_dictionary)
    print(image_data_dictionary.keys())

    # this is to get the title of the image
    image_title_string = image_data_dictionary['title']
    park_word_document.add_paragraph(image_title_string)

    # this is to get the altext
    image_altText_string = image_data_dictionary['altText']
    park_word_document.add_paragraph(image_altText_string)

    # this is to get the image credit
    image_credit_string = image_data_dictionary['credit']
    park_word_document.add_paragraph(image_credit_string)

    # this is to get the caption of the image
    image_caption_string = image_data_dictionary['caption']
    park_word_document.add_paragraph(image_caption_string)

    # this is to get the park picture title
    image_title_string = image_data_dictionary['title']
    park_word_document.add_paragraph(image_title_string)

    # This is to get the image url
    image_download_url = image_data_dictionary['url']
    print(image_download_url)
    park_word_document.add_paragraph(image_download_url)

    # this code is to get images from specific parks
    # we also download the images for each park
    park_image_response = requests.get(image_download_url)

    # the wb is used to ope the file in binary format for writing
    # the with statement is used to open and write to a file
    with open('park_picture.jpg', 'wb') as park_image_file:
        for chunk in park_image_response.iter_content():
            park_image_file.write(chunk)

    # add_picture is used to add an image to a Word document
    # we can also add the picture width and height using docx.shared.Inches
    # you can either write it inline or you can import the import shared at the
    # top of the page and just indicate the number of inches. It is always in Inches or cm
    park_word_document.add_picture('park_picture.jpg', width=docx.shared.Inches(6))

# 6. Save word document.
# we are saving the Word document by using document.save keyword.
park_word_document.save('National Park Guide.docx')
# park_guide.save('ParkGuide.docx')
