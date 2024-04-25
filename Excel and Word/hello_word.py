import docx

# ALWAYS USE THIS LINE AFTER importing docx
document = docx.Document()
# here you are adding the text hello world to the doc
document.add_paragraph('Hello Word', 'Title')

# to add a paragraph
document.add_paragraph('By Eyabezoh Ndi Boniface', 'Heading 1')

# to add text under your paragraph
document.add_paragraph('This is a word document created with python and python-docx')


document.add_paragraph('Automate the boring stuff.', 'Quote')

document.add_paragraph('list of favourite colors', 'Heading 2')

favourite_colors = ['Blue', 'Purple', 'Black']
# using a loop to write the colors to a word document
for color in favourite_colors:
    document.add_paragraph(color, 'List Bullet')
#     List bullet is the style

# here you are saving the document
# when you run it u will see the text
document.save('hello_word.docx')