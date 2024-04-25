import docx
# always import docx
# add the line below too
document = docx.Document()

# list of cars
electric_cars = {
    'Chevrolet': 'Volt',
    'Nissan': 'Leaf',
    'Ford': 'F1'
    'Dodge' 'Mustang',
    'Tesla': 'Model Y'
}

# the title is electric cars and style is heading 1
document.add_paragraph('Electric Cars', 'Heading 1')

# a for loop to write the evs into a Word doc.
for make, model in electric_cars.items():
    document.add_paragraph(make + 'Heading 3')
    # this will append both the mke and model into the word doc
    document.add_paragraph(f'An electric car by {make} is {model}')

# saving the document
document.save('electric_cars.docx')
